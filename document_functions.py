# document_functions.py

import os
from docx import Document
import chromadb
from chromadb.config import Settings
from openai import OpenAI
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

def create_word_doc_from_json(data, filename='output.docx'):
    """Creates a Word document with data and enhanced response."""
    doc = Document()
    doc.add_heading('Generated Research Report', 0)

    # Research details
    doc.add_paragraph(f"Research Topic: {data['research_topic']}")
    doc.add_paragraph(f"Field of Study: {data.get('field_of_study', 'Not Specified')}")
    doc.add_paragraph(f"Publication Type: {data.get('type_of_publication', 'Not Specified')}")

    # Generated response section
    doc.add_heading('Research Summary', level=1)
    doc.add_paragraph(data['response'])

    # Add articles via chosen citation format
    doc.add_heading('Related Articles', level=1)
    for article in data.get('articles', []):
        if data['citation_format'] == 'APA':
            citation = f"{', '.join(article.get('authors', ['Unknown']))} ({article['published'][:4]}). {article['title']}. Retrieved from {article.get('url', 'N/A')}."
        elif data['citation_format'] == 'MLA':
            citation = f"{', '.join(article.get('authors', ['Unknown']))}. \"{article['title']}\". {article.get('url', 'N/A')}, {article['published'][:4]}."
        doc.add_paragraph(citation)

    # Save the document
    doc.save(filename)
    print(f"Document saved successfully as {filename}")
    return filename