# document_functions.py

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI
from dotenv import load_dotenv
import json

# Load environment variables
load_dotenv()

def create_word_doc_from_json(data, filename='output.docx'):
    """
    Creates a Word document from the provided JSON data.
    
    Args:
        data (dict): Dictionary containing research topic, response, and articles
    """
    try:
        # Create a new Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Generated Research Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add research topic
        topic_heading = doc.add_heading('Research Topic:', level=1)
        doc.add_paragraph(data['research_topic'])
        
        # Add response
        summary_heading = doc.add_heading('Research Summary:', level=1)
        doc.add_paragraph(data['response'])
        
        # Add related articles
        if data.get('articles', []):
            articles_heading = doc.add_heading('Related Articles:', level=1)
            
            for article in data.get('articles', []):
                # Format authors
                authors = article.get('authors', [])
                if not authors:
                    author_text = "No authors listed"
                elif len(authors) == 1:
                    author_text = authors[0]
                elif len(authors) == 2:
                    author_text = f"{authors[0]} & {authors[1]}"
                else:
                    author_text = f"{authors[0]} et al."
                
                # Extract year from publication date
                pub_date = article.get('published', '')
                try:
                    # Handle string representation of dict
                    if isinstance(pub_date, str) and pub_date.startswith('{'):
                        import ast
                        date_dict = ast.literal_eval(pub_date)
                        pub_date = date_dict.get('$', '')
                    pub_year = pub_date.split('-')[0] if pub_date else 'N/A'
                except:
                    pub_year = 'N/A'
                
                # Format DOI as URL or use alternative URL
                doi = article.get('doi', '')
                url = ''
                if doi:
                    url = f"https://doi.org/{doi}"
                    source_text = f"DOI: {url}"
                else:
                    # Check for alternative URLs
                    if article.get('url'):
                        url = article['url']
                        # Determine source based on URL
                        if 'arxiv.org' in url.lower():
                            source_text = f"Retrieved from arXiv: {url}"
                        elif 'semanticscholar.org' in url.lower():
                            source_text = f"Retrieved from Semantic Scholar: {url}"
                        elif 'core.ac.uk' in url.lower():
                            source_text = f"Retrieved from CORE: {url}"
                        else:
                            source_text = f"Retrieved from: {url}"
                    else:
                        source_text = 'No URL available'
                
                # Format journal information
                journal_info = []
                if article.get('journal'):
                    journal_info.append(article['journal'])
                if article.get('volume'):
                    journal_info.append(f"Vol. {article['volume']}")
                if article.get('issue'):
                    journal_info.append(f"No. {article['issue']}")
                if article.get('pages'):
                    journal_info.append(f"pp. {article['pages']}")
                journal_text = ', '.join(filter(None, journal_info))
                
                # Create citation based on format
                if data['citation_format'] == 'APA':
                    if journal_text:
                        citation = f"{author_text} ({pub_year}). {article['title']}. {journal_text}. {source_text}"
                    else:
                        citation = f"{author_text} ({pub_year}). {article['title']}. {source_text}"
                elif data['citation_format'] == 'MLA':
                    if journal_text:
                        citation = f"{author_text}. \"{article['title']}\". {journal_text}, {pub_year}. {source_text}"
                    else:
                        citation = f"{author_text}. \"{article['title']}\". {source_text}, {pub_year}."
                
                doc.add_paragraph(citation)
        
        # Save the document
        doc.save(filename)
        print(f"Document saved successfully as {filename}")
        return filename
    except Exception as e:
        print(f"Error creating document: {str(e)}")
        return False

def create_research_summary(data):
    """
    Creates a research summary with title and citations from the provided JSON data.
    
    Args:
        data (dict): Dictionary containing research topic, response, and articles
    
    Returns:
        dict: Dictionary containing research title, summary, and formatted citations
    """
    try:
        summary = {
            'title': 'Generated Research Report',
            'research_topic': data['research_topic'],
            'research_summary': data['response'],
            'citations': []
        }
        
        # Process articles and create citations
        if data.get('articles', []):
            for article in data.get('articles', []):
                # Format authors
                authors = article.get('authors', [])
                if not authors:
                    author_text = "No authors listed"
                elif len(authors) == 1:
                    author_text = authors[0]
                elif len(authors) == 2:
                    author_text = f"{authors[0]} & {authors[1]}"
                else:
                    author_text = f"{authors[0]} et al."
                
                # Extract year from publication date
                pub_date = article.get('published', '')
                try:
                    if isinstance(pub_date, str) and pub_date.startswith('{'):
                        import ast
                        date_dict = ast.literal_eval(pub_date)
                        pub_date = date_dict.get('$', '')
                    pub_year = pub_date.split('-')[0] if pub_date else 'N/A'
                except:
                    pub_year = 'N/A'
                
                # Format DOI as URL or use alternative URL
                doi = article.get('doi', '')
                url = ''
                if doi:
                    url = f"https://doi.org/{doi}"
                    source_text = f"DOI: {url}"
                else:
                    if article.get('url'):
                        url = article['url']
                        if 'arxiv.org' in url.lower():
                            source_text = f"Retrieved from arXiv: {url}"
                        elif 'semanticscholar.org' in url.lower():
                            source_text = f"Retrieved from Semantic Scholar: {url}"
                        elif 'core.ac.uk' in url.lower():
                            source_text = f"Retrieved from CORE: {url}"
                        else:
                            source_text = f"Retrieved from: {url}"
                    else:
                        source_text = 'No URL available'
                
                # Format journal information
                journal_info = []
                if article.get('journal'):
                    journal_info.append(article['journal'])
                if article.get('volume'):
                    journal_info.append(f"Vol. {article['volume']}")
                if article.get('issue'):
                    journal_info.append(f"No. {article['issue']}")
                if article.get('pages'):
                    journal_info.append(f"pp. {article['pages']}")
                journal_text = ', '.join(filter(None, journal_info))
                
                # Create citation based on format
                citation = {}
                citation['title'] = article['title']
                citation['authors'] = author_text
                citation['year'] = pub_year
                citation['url'] = url
                citation['journal_info'] = journal_text if journal_text else None
                
                if data['citation_format'] == 'APA':
                    if journal_text:
                        citation['formatted'] = f"{author_text} ({pub_year}). {article['title']}. {journal_text}. {source_text}"
                    else:
                        citation['formatted'] = f"{author_text} ({pub_year}). {article['title']}. {source_text}"
                elif data['citation_format'] == 'MLA':
                    if journal_text:
                        citation['formatted'] = f"{author_text}. \"{article['title']}\". {journal_text}, {pub_year}. {source_text}"
                    else:
                        citation['formatted'] = f"{author_text}. \"{article['title']}\". {source_text}, {pub_year}."
                
                summary['citations'].append(citation)
        
        return summary
    except Exception as e:
        print(f"Error creating summary: {str(e)}")
        return None